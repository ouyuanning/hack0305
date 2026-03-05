#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成功能迭代介绍 Word 文档（一页内），供团队成员了解近期更新。
输出到 ~/Downloads 目录。
"""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    return p


def add_para(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
    return p


def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(10.5)

    # 页面边距缩小，便于一页内展示
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 标题
    title = doc.add_paragraph()
    t = title.add_run('GitHub Issue 智能管理系统 · 功能迭代介绍')
    t.bold = True
    t.font.size = Pt(16)
    t.font.name = 'Microsoft YaHei'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(4)
    add_para(doc, '版本 V1 | 2026-02 | 供团队成员了解近期更新', indent=False)

    add_heading(doc, '一、本次更新要点', 1)
    add_para(doc, '• 提 Bug / Feature：自然语言描述 → AI 生成 Issue 草稿 → 本地 GitHub 风格预览 → 一键创建')
    add_para(doc, '• 智能上下文检测：CDP / 窗口标题 / 剪贴板 / 手动输入，4 层降级')
    add_para(doc, '• 模板自动推断：MO Bug、MOI Bug、User Bug、Doc Request 等按关键词或浏览器 Issue 标签')
    add_para(doc, '• 项目看板：按 project/xxx 同步 Issue → 每日生成进度/逾期/阻塞看板')

    add_heading(doc, '二、核心功能', 1)
    add_para(doc, '提 Issue 模块：输入一句话描述（如「moi 跨页表识别内容错误」），系统自动推断类型、填充模板（环境/复现/截图等）、推荐标签和负责人；支持 --preview 仅生成本地网页不提交 GitHub。')
    add_para(doc, '项目看板：sync_project_issues.py 同步带项目标签的 Issue，generate_daily_dashboard.py 生成 Markdown 看板（可含 AI 总结）。')

    add_heading(doc, '三、工作原理', 1)
    add_para(doc, '1）上下文获取：依次尝试 Chrome CDP（需 --remote-debugging-port=9222）、系统窗口标题、剪贴板、手动输入。')
    add_para(doc, '2）类型推断：优先使用浏览器 Issue 标签；若无则 AI + 知识库；再则关键词（mo/moi/问数/docs 等）。')
    add_para(doc, '3）草稿生成：将用户描述 + 上下文 + 模板结构交给 AI（通义千问/Claude），输出 JSON 草稿。')
    add_para(doc, '4）预览与创建：write_preview_html 生成 GitHub 风格两栏预览页；确认后调用 GitHub API 创建 Issue。')

    add_heading(doc, '四、快速使用', 1)
    add_para(doc, '预览：python3 feature_issue_and_kanban/scripts/create_issue_interactive.py --input "描述" --repo matrixorigin/matrixflow --preview --output-html feature_issue_and_kanban/preview.html')
    add_para(doc, '直接创建：去掉 --preview。交互模式：加 --interactive。')
    add_para(doc, '')
    add_para(doc, '—— 文档结束 ——')

    out_dir = Path.home() / 'Downloads'
    out_path = out_dir / 'GitHub_Issue智能管理系统_功能迭代介绍-260224.docx'
    doc.save(str(out_path))
    print(f"已生成: {out_path}")
    return str(out_path)


if __name__ == '__main__':
    main()
